# =============================================================================
# dbv-md2word — Conversor personalizable de Markdown a Word (.docx) con interfaz visual local
# Copyright (c) 2026 David Bueno Vallejo · https://github.com/davidbuenov
# Licensed under the MIT License. See LICENSE for details.
# Built with dbv-specs-ops · https://github.com/davidbuenov/dbv-specs-ops
#
# PARCHE (ver CHANGELOG al final del fichero):
#   1. Fix: falso positivo en la detección de enlaces markdown dentro de
#      add_runs_to_paragraph (celdas tipo "[Sí] Activo (pasivo)" perdían texto).
#   2. Feature: los emoji (✅ ❌ ⚠️ 🟡 🟢 🧟 etc.) ahora se renderizan con una
#      fuente específica (por defecto "Segoe UI Emoji") en lugar de heredar
#      la fuente del cuerpo (Aptos), que no tiene esos glifos.
# =============================================================================

import re
import os
import sys
import copy
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.opc.constants import RELATIONSHIP_TYPE

__version__ = "1.4.1"

class ConversionState:
    """Clase para mantener el estado de la conversión, como los contadores de marcadores."""
    def __init__(self):
        self.bookmark_counter = 0

    def get_next_bookmark_id(self):
        self.bookmark_counter += 1
        return self.bookmark_counter

def get_style_safely(doc, english_name, spanish_name):
    """Busca un estilo en el documento de forma segura por su nombre en inglés o español."""
    for name in [english_name, spanish_name]:
        if name in doc.styles:
            return doc.styles[name]
    # Intenta buscar por ID (eliminando espacios)
    style_id = english_name.replace(" ", "")
    for s in doc.styles:
        if s.style_id == style_id:
            return s
    return None

def get_or_create_paragraph_style(doc, name, base_style_name='Normal'):
    """Retorna un estilo de párrafo o lo crea si no existe y lo añade a la galería."""
    try:
        return doc.styles[name]
    except KeyError:
        from docx.enum.style import WD_STYLE_TYPE
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles[base_style_name]
        style.quick_style = True
        return style

def get_or_create_character_style(doc, name, base_style_name='Default Paragraph Font'):
    """Retorna un estilo de carácter o lo crea si no existe y lo añade a la galería."""
    try:
        return doc.styles[name]
    except KeyError:
        from docx.enum.style import WD_STYLE_TYPE
        style = doc.styles.add_style(name, WD_STYLE_TYPE.CHARACTER)
        try:
            style.base_style = doc.styles[base_style_name]
        except KeyError:
            pass
        style.quick_style = True
        return style

def sanitize_bookmark_name(text):
    """Sanitiza el texto para crear un nombre de marcador válido en Word."""
    sanitized = re.sub(r'[^a-zA-Z0-9_]', '', text)
    if not sanitized:
        sanitized = "bookmark"
    if sanitized[0].isdigit():
        sanitized = "_" + sanitized
    return sanitized[:40]

def configure_document_styles(doc, config):
    """Configura las fuentes y colores de los estilos del documento según la configuración del usuario."""
    # Estilo Normal (Cuerpo del texto)
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = config.get('body_font', 'Calibri')
    font_normal.size = Pt(11)

    body_align = config.get('body_align', 'justify')
    if body_align == 'justify':
        style_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        style_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    primary_color_hex = config.get('primary_color', '1F4E79')
    if primary_color_hex.startswith('#'):
        primary_color_hex = primary_color_hex[1:]
    primary_rgb = RGBColor.from_string(primary_color_hex)

    # Configuración de estilo "Title / Título" para título principal
    style_title = get_style_safely(doc, 'Title', 'Título')
    if style_title:
        font_title = style_title.font
        font_title.name = config.get('heading_font', 'Aptos Display')
        font_title.size = Pt(24)
        font_title.bold = True
        font_title.color.rgb = primary_rgb
        style_title.quick_style = True

    # Configuración de Encabezados Nativos
    heading_configs = [
        (1, 'Heading 1', 'Título 1', Pt(20), True),
        (2, 'Heading 2', 'Título 2', Pt(15), True),
        (3, 'Heading 3', 'Título 3', Pt(12.5), True),
        (4, 'Heading 4', 'Título 4', Pt(11), True),
    ]

    for lvl, eng_name, esp_name, size, bold in heading_configs:
        style = get_style_safely(doc, eng_name, esp_name)
        if style:
            font = style.font
            font.name = config.get('heading_font', 'Aptos Display')
            font.size = size
            font.bold = bold
            font.color.rgb = primary_rgb
            style.quick_style = True

            # Márgenes de párrafo para encabezados
            style.paragraph_format.keep_with_next = True
            if lvl == 1:
                style.paragraph_format.space_before = Pt(20)
                style.paragraph_format.space_after = Pt(8)
            elif lvl == 2:
                style.paragraph_format.space_before = Pt(14)
                style.paragraph_format.space_after = Pt(6)
            else:
                style.paragraph_format.space_before = Pt(10)
                style.paragraph_format.space_after = Pt(4)

    # Estilo de bloque de código ('codigo' - párrafo)
    style_code = get_or_create_paragraph_style(doc, 'codigo')
    font_code = style_code.font
    font_code.name = config.get('code_font', 'Consolas')
    font_code.size = Pt(9.5)
    font_code.color.rgb = RGBColor(60, 60, 60)
    style_code.paragraph_format.space_after = Pt(0)
    style_code.paragraph_format.line_spacing = 1.0
    style_code.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Forzar alineación a la izquierda
    style_code.quick_style = True

    # Estilo de código en línea ('codigo_car' - carácter)
    style_code_car = get_or_create_character_style(doc, 'codigo_car')
    font_code_car = style_code_car.font
    font_code_car.name = config.get('code_font', 'Consolas')
    font_code_car.size = Pt(9.5)
    style_code_car.quick_style = True

    # XML para sombreado de fondo gris en inline code
    rPr = style_code_car.element.get_or_add_rPr()
    # Limpiamos sombreados previos si existen
    for child in list(rPr):
        if child.tag.endswith('shd'):
            rPr.remove(child)
    shading = parse_xml(r'<w:shd {} w:fill="F0F0F0"/>'.format(nsdecls('w')))
    rPr.append(shading)

def add_field(paragraph, field_code, placeholder_text, font_name, is_bold=False, color_rgb=None):
    """Agrega un campo de Word estructurado correctamente (begin -> instrText -> separate -> result -> end)."""
    # 1. Run de inicio del campo
    run_begin = paragraph.add_run()
    run_begin._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))

    # 2. Run con el código del campo
    run_instr = paragraph.add_run()
    run_instr._r.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve">{field_code}</w:instrText>'))

    # 3. Run de separación
    run_sep = paragraph.add_run()
    run_sep._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>'))

    # 4. Run con el resultado/placeholder (va entre separate y end)
    run_val = paragraph.add_run(placeholder_text)
    run_val.font.name = font_name
    if is_bold:
        run_val.bold = True
    if color_rgb:
        run_val.font.color.rgb = color_rgb

    # 5. Run de fin de campo
    run_end = paragraph.add_run()
    run_end._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))

def add_toc(doc, config):
    """Inserta una Tabla de Contenidos nativa de Word."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)

    primary_color_hex = config.get('primary_color', '1F4E79')
    if primary_color_hex.startswith('#'):
        primary_color_hex = primary_color_hex[1:]
    primary_rgb = RGBColor.from_string(primary_color_hex)

    run_title = p.add_run("Índice de Contenidos")
    run_title.bold = True
    run_title.font.name = config.get('heading_font', 'Aptos Display')
    run_title.font.size = Pt(14)
    run_title.font.color.rgb = primary_rgb

    p_desc = doc.add_paragraph()
    p_desc.paragraph_format.space_after = Pt(4)
    run_desc = p_desc.add_run("[Haz clic derecho sobre el bloque de abajo y selecciona 'Actualizar campos' para actualizar el índice]")
    run_desc.font.name = config.get('body_font', 'Aptos')
    run_desc.font.italic = True
    run_desc.font.size = Pt(9.5)
    run_desc.font.color.rgb = RGBColor(120, 120, 120)

    p_toc = doc.add_paragraph()
    p_toc.paragraph_format.space_after = Pt(12)

    body_font = config.get('body_font', 'Aptos')
    add_field(
        p_toc,
        field_code=' TOC \\o "1-3" \\h \\z \\u ',
        placeholder_text="--- Índice Generado por Word ---",
        font_name=body_font,
        color_rgb=RGBColor(128, 128, 128)
    )

def add_figure_caption(doc, text, bookmark_name, state, config):
    """Añade un pie de foto de figura con numeración automática SEQ y marcador para referencias."""
    p = doc.add_paragraph(style='Caption')
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(16)
    p.paragraph_format.keep_with_next = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    body_font = config.get('body_font', 'Aptos')
    run_prefix = p.add_run("Fig. ")
    run_prefix.font.name = body_font

    # Campo SEQ
    add_field(
        p,
        field_code=' SEQ Figure \\* ARABIC ',
        placeholder_text="0",
        font_name=body_font,
        is_bold=True
    )

    if text:
        run_text = p.add_run(f": {text}")
        run_text.font.name = body_font

    # Inyectar Marcador (Bookmark)
    bookmark_id = state.get_next_bookmark_id()
    start = parse_xml(f'<w:bookmarkStart {nsdecls("w")} w:id="{bookmark_id}" w:name="{bookmark_name}"/>')
    end = parse_xml(f'<w:bookmarkEnd {nsdecls("w")} w:id="{bookmark_id}"/>')
    p._p.insert(0, start)
    p._p.append(end)

def add_table_caption(doc, text, bookmark_name, state, config):
    """Añade un título de tabla superior con numeración automática SEQ y marcador."""
    p = doc.add_paragraph(style='Caption')
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Forzar alineación a la izquierda

    body_font = config.get('body_font', 'Aptos')
    run_prefix = p.add_run("Tabla ")
    run_prefix.font.name = body_font

    # Campo SEQ
    add_field(
        p,
        field_code=' SEQ Table \\* ARABIC ',
        placeholder_text="0",
        font_name=body_font,
        is_bold=True
    )

    if text:
        run_text = p.add_run(f": {text}")
        run_text.font.name = body_font

    # Inyectar Marcador (Bookmark)
    bookmark_id = state.get_next_bookmark_id()
    start = parse_xml(f'<w:bookmarkStart {nsdecls("w")} w:id="{bookmark_id}" w:name="{bookmark_name}"/>')
    end = parse_xml(f'<w:bookmarkEnd {nsdecls("w")} w:id="{bookmark_id}"/>')
    p._p.insert(0, start)
    p._p.append(end)

def add_cross_reference(paragraph, bookmark_name, prefix_text, body_font='Aptos'):
    """Agrega un campo de referencia cruzada REF dinámico en el párrafo."""
    r_prefix = paragraph.add_run(prefix_text)
    r_prefix.font.name = body_font

    add_field(
        paragraph,
        field_code=f' REF {bookmark_name} \\h ',
        placeholder_text="?",
        font_name=body_font,
        is_bold=True
    )

def add_hyperlink(paragraph, text, url, body_font='Aptos'):
    """Añade un hipervínculo en color azul y subrayado con la fuente del cuerpo."""
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0563C1')
    rPr.append(c)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    # Inyectamos fuente para hipervínculo
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), body_font)
    rFonts.set(qn('w:hAnsi'), body_font)
    rPr.append(rFonts)

    new_run.append(rPr)

    text_node = OxmlElement('w:t')
    text_node.text = text
    new_run.append(text_node)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def add_runs_to_paragraph(p, text, config, state):
    """Parsea el texto markdown y añade runs con estilos y referencias cruzadas."""
    # Buscamos negrita, cursiva, código, links e indicaciones de referencias cruzadas
    pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*|`.*?`|\[.*?\]\(.*?\)|\[(?:Fig\.|Figura\s+|Tabla\s+|Tabla\.)\s*.*?\])')
    parts = pattern.split(text)

    body_font = config.get('body_font', 'Calibri')
    code_font = config.get('code_font', 'Consolas')

    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2])
            r.bold = True
            r.font.name = body_font
        elif part.startswith('*') and part.endswith('*'):
            r = p.add_run(part[1:-1])
            r.italic = True
            r.font.name = body_font
        elif part.startswith('`') and part.endswith('`'):
            r = p.add_run(part[1:-1])
            r.style = 'codigo_car'
            r.font.name = code_font
        # --- PARCHE (fix #1): antes se comprobaba '[' ... ']' ... '(' ... ')' en
        # cualquier posición del texto, lo que confundía celdas como
        # "[Sí] Activo (pasivo)" con un enlace markdown y perdía el resto del
        # texto. Ahora exigimos la adyacencia real "](" propia de [texto](url).
        elif part.startswith('[') and '](' in part and part.endswith(')'):
            link_match = re.match(r'^\[(.*?)\]\((.*?)\)$', part)
            if link_match:
                link_text = link_match.group(1)
                link_url = link_match.group(2)
                add_hyperlink(p, link_text, link_url, body_font)
            else:
                r = p.add_run(part)
                r.font.name = body_font
        elif part.startswith('[') and part.endswith(']'):
            # Posible referencia cruzada
            match_ref = re.match(r'^\[(Fig\.|Figura\s+|Tabla\s+|Tabla\.)\s*(.*?)\]$', part, re.IGNORECASE)
            if match_ref:
                ref_type = match_ref.group(1).strip().lower()
                label = match_ref.group(2).strip()
                if 'fig' in ref_type:
                    bookmark_name = f"_Ref_Fig_{sanitize_bookmark_name(label)}"
                    add_cross_reference(p, bookmark_name, "Fig. ", body_font)
                else:
                    bookmark_name = f"_Ref_Tabla_{sanitize_bookmark_name(label)}"
                    add_cross_reference(p, bookmark_name, "Tabla ", body_font)
            else:
                r = p.add_run(part)
                r.font.name = body_font
        else:
            r = p.add_run(part)
            r.font.name = body_font

def set_cell_background(cell, color_hex):
    """Establece el fondo de una celda de tabla."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    tcPr.append(shd)

# =============================================================================
# PARCHE (feature #2): fuente dedicada para emoji.
#
# Word solo renderiza en color los emoji cuando la fuente del carácter
# soporta esos glifos (p. ej. "Segoe UI Emoji" en Windows/Office). Si se deja
# la fuente del cuerpo (Aptos, Calibri, etc.) tal cual, el emoji aparece como
# un cuadro vacío ("tofu"). Estas funciones detectan los caracteres Unicode
# de tipo emoji dentro de cada run ya generado y los aíslan en su propio
# sub-run con la fuente de emoji, sin tocar el resto del texto.
# =============================================================================

EMOJI_RANGES = [
    (0x2190, 0x21FF),   # Flechas
    (0x2300, 0x23FF),   # Misceláneos técnicos (⏳ etc.)
    (0x2460, 0x24FF),   # Encerrados alfanuméricos
    (0x25A0, 0x25FF),   # Formas geométricas (▲ ▼ ● etc.)
    (0x2600, 0x27BF),   # Símbolos misceláneos y Dingbats (✅ ❌ ⚠ ☎ etc.)
    (0x2B00, 0x2BFF),   # Flechas y símbolos adicionales
    (0xFE00, 0xFE0F),   # Selectores de variación (VS1-VS16)
    (0x1F000, 0x1FAFF), # Bloques de emoji suplementarios (🟡 🟢 🧟 etc.)
]

def is_emoji_char(ch):
    """Determina si un carácter cae dentro de los rangos Unicode típicos de emoji."""
    cp = ord(ch)
    for start, end in EMOJI_RANGES:
        if start <= cp <= end:
            return True
    return False

def _split_run_by_emoji(run, emoji_font):
    """Divide un run existente en sub-runs (texto normal / emoji), aplicando
    emoji_font únicamente a los caracteres emoji y conservando el resto del
    formato (negrita, cursiva, color, fuente original) del run de origen."""
    text = run.text
    if not text or not any(is_emoji_char(ch) for ch in text):
        return

    segments = []
    current, current_is_emoji = "", None
    for ch in text:
        ch_is_emoji = is_emoji_char(ch)
        if current_is_emoji is None:
            current, current_is_emoji = ch, ch_is_emoji
        elif ch_is_emoji == current_is_emoji:
            current += ch
        else:
            segments.append((current, current_is_emoji))
            current, current_is_emoji = ch, ch_is_emoji
    if current:
        segments.append((current, current_is_emoji))

    r_element = run._element
    rPr_template = r_element.find(qn('w:rPr'))

    new_elements = []
    for seg_text, seg_is_emoji in segments:
        new_r = OxmlElement('w:r')
        new_rPr = copy.deepcopy(rPr_template) if rPr_template is not None else OxmlElement('w:rPr')
        if seg_is_emoji:
            existing_rFonts = new_rPr.find(qn('w:rFonts'))
            if existing_rFonts is not None:
                new_rPr.remove(existing_rFonts)
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), emoji_font)
            rFonts.set(qn('w:hAnsi'), emoji_font)
            rFonts.set(qn('w:cs'), emoji_font)
            new_rPr.insert(0, rFonts)
        new_r.append(new_rPr)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = seg_text
        new_r.append(t)
        new_elements.append(new_r)

    parent = r_element.getparent()
    idx = list(parent).index(r_element)
    for i, el in enumerate(new_elements):
        parent.insert(idx + i, el)
    parent.remove(r_element)

def apply_emoji_font(doc, emoji_font):
    """Recorre todo el documento (párrafos sueltos y celdas de tabla, incluidas
    tablas anidadas) aplicando la fuente de emoji donde corresponda."""
    if not emoji_font:
        return

    def process_paragraph(p):
        for run in list(p.runs):
            _split_run_by_emoji(run, emoji_font)

    def process_table(table):
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    process_paragraph(p)
                for nested_table in cell.tables:
                    process_table(nested_table)

    for p in doc.paragraphs:
        process_paragraph(p)
    for table in doc.tables:
        process_table(table)

def parse_markdown(filepath):
    """Parsea el archivo Markdown y genera una lista de bloques estructurados."""
    if not os.path.exists(filepath):
        print(f"Error: El archivo {filepath} no existe.")
        sys.exit(1)

    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    blocks = []
    current_block = None
    in_code = False
    in_table = False
    table_caption_pending = None

    for line in lines:
        stripped = line.strip()

        # Manejo de bloques de código
        if stripped.startswith('```'):
            if in_code:
                in_code = False
                blocks.append(current_block)
                current_block = None
            else:
                if current_block:
                    blocks.append(current_block)
                in_code = True
                lang = stripped[3:].strip()
                current_block = {'type': 'code', 'lang': lang, 'lines': []}
            continue

        if in_code:
            current_block['lines'].append(line.rstrip('\n'))
            continue

        # Detección de subtítulo/caption de tablas ("Tabla: Titulo" o "Table: Titulo")
        caption_match = re.match(r'^(Table|Tabla):\s*(.*)$', stripped, re.IGNORECASE)
        if caption_match:
            if current_block:
                blocks.append(current_block)
            table_caption_pending = caption_match.group(2).strip()
            current_block = None
            continue

        # Manejo de tablas
        if stripped.startswith('|'):
            if re.match(r'^\|[\s:\-|]+$', stripped):
                continue
            row_cells = [cell.strip() for cell in stripped.split('|')[1:-1]]
            if in_table:
                current_block['rows'].append(row_cells)
            else:
                if current_block:
                    blocks.append(current_block)
                in_table = True
                current_block = {
                    'type': 'table',
                    'rows': [row_cells],
                    'caption': table_caption_pending
                }
                table_caption_pending = None
            continue
        else:
            if in_table:
                in_table = False
                blocks.append(current_block)
                current_block = None
            elif table_caption_pending:
                if current_block:
                    blocks.append(current_block)
                current_block = {'type': 'paragraph', 'text': f"Tabla: {table_caption_pending}"}
                blocks.append(current_block)
                current_block = None
                table_caption_pending = None

        # Línea vacía
        if not stripped:
            if current_block and current_block['type'] not in ['code', 'table']:
                blocks.append(current_block)
                current_block = None
            continue

        # Detección de imágenes (Figura)
        img_match = re.match(r'^!\[(.*?)\]\((.*?)\)$', stripped)
        if img_match:
            if current_block:
                blocks.append(current_block)
            alt = img_match.group(1).strip()
            src = img_match.group(2).strip()
            blocks.append({'type': 'image', 'alt': alt, 'src': src})
            current_block = None
            continue

        # Encabezados (H1 - H6)
        heading_match = re.match(r'^(#{1,6})\s+(.*)$', stripped)
        if heading_match:
            if current_block:
                blocks.append(current_block)
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            blocks.append({'type': 'heading', 'level': level, 'text': text})
            current_block = None
            continue

        # Listas con viñetas o numeradas
        bullet_match = re.match(r'^([-*]|\d+\.)\s+(.*)$', stripped)
        indent_match = re.match(r'^(\s+)([-*]|\d+\.)\s+(.*)$', line)

        if indent_match:
            indent_spaces = len(indent_match.group(1))
            text = indent_match.group(3)
            level = indent_spaces // 2
            item = {'text': text, 'level': level}
            if current_block and current_block['type'] == 'list':
                current_block['items'].append(item)
            else:
                if current_block:
                    blocks.append(current_block)
                current_block = {'type': 'list', 'items': [item]}
            continue
        elif bullet_match:
            text = bullet_match.group(2)
            item = {'text': text, 'level': 0}
            if current_block and current_block['type'] == 'list':
                current_block['items'].append(item)
            else:
                if current_block:
                    blocks.append(current_block)
                current_block = {'type': 'list', 'items': [item]}
            continue

        # Línea horizontal
        if stripped == '---':
            if current_block:
                blocks.append(current_block)
            blocks.append({'type': 'hr'})
            current_block = None
            continue

        # Párrafos normales
        if current_block and current_block['type'] == 'paragraph':
            current_block['text'] += " " + stripped
        else:
            if current_block:
                blocks.append(current_block)
            current_block = {'type': 'paragraph', 'text': stripped}

    if current_block:
        blocks.append(current_block)

    return blocks

def create_docx(blocks, output_path, config=None, base_dir=None):
    """Genera un archivo .docx a partir de bloques estructurados y una configuración."""
    if config is None:
        config = {}

    doc = Document()
    state = ConversionState()

    # Configurar márgenes estándar (1 pulgada = 2.54 cm)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Configurar estilos de fuentes y colores
    configure_document_styles(doc, config)

    toc_added = False
    toc_enabled = config.get('toc_enabled', True)
    shift_headings = config.get('shift_headings', False)

    heading_font = config.get('heading_font', 'Aptos Display')
    body_font = config.get('body_font', 'Aptos')

    primary_color_hex = config.get('primary_color', '1F4E79')
    if primary_color_hex.startswith('#'):
        primary_color_hex = primary_color_hex[1:]
    primary_rgb = RGBColor.from_string(primary_color_hex)

    for block in blocks:
        # Si la Tabla de Contenidos está habilitada y no ha sido agregada,
        # la agregamos al inicio a menos que este bloque sea el título principal (nivel 1).
        if toc_enabled and not toc_added:
            is_main_title = False
            if block['type'] == 'heading':
                if shift_headings and block['level'] == 1:
                    is_main_title = True
                elif not shift_headings and block['level'] == 1:
                    is_main_title = True

            if not is_main_title:
                add_toc(doc, config)
                toc_added = True
        if block['type'] == 'heading':
            level = block['level']
            text = block['text']

            p = doc.add_paragraph()
            p.paragraph_format.keep_with_next = True

            # Caso 1: Desplazamiento habilitado y es H1 (se convierte en Título Principal del documento)
            if level == 1 and shift_headings:
                p.style = get_style_safely(doc, 'Title', 'Título') or 'Title'
                run = p.add_run(text)
                run.font.name = heading_font
                run.font.size = Pt(24)
                run.bold = True
                run.font.color.rgb = primary_rgb

                # Inyectamos el TOC tras el Título Principal si está habilitado
                if toc_enabled and not toc_added:
                    add_toc(doc, config)
                    toc_added = True
            else:
                # Si está desplazado, restamos 1 al nivel (H2 -> Título 1, etc.)
                actual_level = level - 1 if shift_headings else level
                if actual_level <= 0:
                    actual_level = 1

                if actual_level == 1:
                    p.style = get_style_safely(doc, 'Heading 1', 'Título 1') or 'Heading 1'
                    run = p.add_run(text)
                    run.font.name = heading_font
                    run.font.size = Pt(20)
                    run.bold = True
                    run.font.color.rgb = primary_rgb

                    if toc_enabled and not toc_added:
                        add_toc(doc, config)
                        toc_added = True
                elif actual_level == 2:
                    p.style = get_style_safely(doc, 'Heading 2', 'Título 2') or 'Heading 2'
                    run = p.add_run(text)
                    run.font.name = heading_font
                    run.font.size = Pt(15)
                    run.bold = True
                    run.font.color.rgb = primary_rgb
                elif actual_level == 3:
                    p.style = get_style_safely(doc, 'Heading 3', 'Título 3') or 'Heading 3'
                    run = p.add_run(text)
                    run.font.name = heading_font
                    run.font.size = Pt(12.5)
                    run.bold = True
                    run.font.color.rgb = primary_rgb
                else:
                    p.style = get_style_safely(doc, 'Heading 4', 'Título 4') or 'Heading 4'
                    run = p.add_run(text)
                    run.font.name = heading_font
                    run.font.size = Pt(11)
                    run.bold = True
                    run.font.color.rgb = primary_rgb

        elif block['type'] == 'paragraph':
            if not block['text'].strip():
                continue
            p = doc.add_paragraph()
            p.style = 'Normal'
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.15
            add_runs_to_paragraph(p, block['text'], config, state)

        elif block['type'] == 'list':
            for item in block['items']:
                style_name = 'List Bullet'
                if item['level'] == 1:
                    style_name = 'List Bullet 2'
                elif item['level'] >= 2:
                    style_name = 'List Bullet 3'

                p = doc.add_paragraph(style=style_name)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.line_spacing = 1.15
                add_runs_to_paragraph(p, item['text'], config, state)

        elif block['type'] == 'hr':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            p_border = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:bottom w:val="single" w:sz="6" w:space="1" w:color="D3D3D3"/></w:pBdr>')
            p._p.get_or_add_pPr().append(p_border)

        elif block['type'] == 'code':
            # Caja para el bloque de código con tabla 1x1
            table = doc.add_table(rows=1, cols=1)
            table.autofit = False
            cell = table.cell(0, 0)
            cell.width = Inches(6.5)

            # Sombreado de fondo
            code_bg = config.get('code_bg_color', 'F5F5F5')
            if code_bg.startswith('#'):
                code_bg = code_bg[1:]
            set_cell_background(cell, code_bg)

            # Borde izquierdo con color corporativo
            border_color = config.get('primary_color', '1F4E79')
            if border_color.startswith('#'):
                border_color = border_color[1:]

            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}>'
                                  f'<w:top w:val="none"/>'
                                  f'<w:left w:val="single" w:sz="18" w:space="0" w:color="{border_color}"/>'
                                  f'<w:bottom w:val="none"/>'
                                  f'<w:right w:val="none"/>'
                                  f'</w:tcBorders>')
            tcPr.append(tcBorders)

            p = cell.paragraphs[0]
            p.style = 'codigo'

            code_text = "\n".join(block['lines'])
            run_code = p.add_run(code_text)
            run_code.font.name = config.get('code_font', 'Consolas')

            # Salto después de la tabla
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(0)
            spacer.paragraph_format.space_after = Pt(6)

        elif block['type'] == 'image':
            alt = block['alt']
            src = block['src']

            # Resolver ruta
            img_path = os.path.join(base_dir, src) if base_dir else src
            img_path = img_path.split('?')[0].strip('"\'')

            if os.path.exists(img_path):
                try:
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img.paragraph_format.space_before = Pt(12)
                    p_img.paragraph_format.space_after = Pt(4)
                    p_img.paragraph_format.keep_with_next = True

                    run_img = p_img.add_run()
                    run_img.add_picture(img_path, width=Inches(5.5))

                    # Caption
                    bookmark_name = f"_Ref_Fig_{sanitize_bookmark_name(alt)}"
                    add_figure_caption(doc, alt, bookmark_name, state, config)
                except Exception as e:
                    p_err = doc.add_paragraph()
                    r_err = p_err.add_run(f"[Error al insertar imagen '{src}': {str(e)}]")
                    r_err.font.color.rgb = RGBColor(255, 0, 0)
                    r_err.font.italic = True
            else:
                p_err = doc.add_paragraph()
                r_err = p_err.add_run(f"[Imagen no encontrada: '{src}']")
                r_err.font.color.rgb = RGBColor(255, 0, 0)
                r_err.font.italic = True

        elif block['type'] == 'table':
            rows_data = block['rows']
            caption = block.get('caption')

            if not rows_data:
                continue

            # Caption arriba de la tabla
            if caption:
                bookmark_name = f"_Ref_Tabla_{sanitize_bookmark_name(caption)}"
                add_table_caption(doc, caption, bookmark_name, state, config)

            cols_count = len(rows_data[0])
            table = doc.add_table(rows=len(rows_data), cols=cols_count)
            table.style = 'Table Grid'
            table.autofit = True

            primary_color_hex = config.get('primary_color', '1F4E79')
            if primary_color_hex.startswith('#'):
                primary_color_hex = primary_color_hex[1:]

            for r_idx, row_data in enumerate(rows_data):
                row = table.rows[r_idx]
                is_header = (r_idx == 0)

                for c_idx, cell_value in enumerate(row_data):
                    if c_idx >= len(row.cells):
                        continue
                    cell = row.cells[c_idx]
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.space_before = Pt(2)

                    if is_header:
                        set_cell_background(cell, primary_color_hex)
                        p.style = 'Normal'
                        add_runs_to_paragraph(p, cell_value, config, state)
                        # Cabeceras en blanco y negrita
                        for r in p.runs:
                            r.bold = True
                            r.font.color.rgb = RGBColor(255, 255, 255)
                    else:
                        if r_idx % 2 == 0:
                            set_cell_background(cell, "F2F5F8")
                        add_runs_to_paragraph(p, cell_value, config, state)

            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(0)
            spacer.paragraph_format.space_after = Pt(8)

    # --- PARCHE (feature #2): aplicar fuente de emoji a todo el documento
    # justo antes de guardar, una vez que todos los runs ya existen.
    emoji_font = config.get('emoji_font', 'Segoe UI Emoji')
    apply_emoji_font(doc, emoji_font)

    out_dir = os.path.dirname(output_path)
    if out_dir and not os.path.exists(out_dir):
        os.makedirs(out_dir, exist_ok=True)
    doc.save(output_path)

def load_default_config(custom_path=None):
    """Carga la configuración por defecto desde config.json o retorna fallbacks, con soporte opcional de custom_path."""
    default_config = {
        'heading_font': 'Aptos Display',
        'body_font': 'Aptos',
        'code_font': 'Consolas',
        'emoji_font': 'Segoe UI Emoji',
        'body_align': 'justify',
        'primary_color': '#1F4E79',
        'toc_enabled': True,
        'numbering_enabled': True,
        'shift_headings': False
    }

    paths_to_load = []
    base_config_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'config.json')
    if os.path.exists(base_config_path):
        paths_to_load.append(base_config_path)
    if custom_path and os.path.exists(custom_path):
        paths_to_load.append(custom_path)

    for p in paths_to_load:
        import json
        try:
            with open(p, 'r', encoding='utf-8') as f:
                loaded = json.load(f)
                for k, v in loaded.items():
                    default_config[k] = v
        except Exception:
            pass
    return default_config

if __name__ == '__main__':
    if len(sys.argv) >= 2 and sys.argv[1] in ('--version', '-v'):
        print(f"dbv-md2word v{__version__}, creado por David Bueno Vallejo — libre y gratuito · https://github.com/davidbuenov")
        sys.exit(0)

    if len(sys.argv) < 2:
        print("Uso: python convert_md_to_docx.py <archivo.md> [archivo.docx] [config.json]")
        print("Uso: python convert_md_to_docx.py --version | -v")
        sys.exit(1)

    src = sys.argv[1]
    dst = None
    config_path = None

    if len(sys.argv) >= 3:
        dst = sys.argv[2]

    if len(sys.argv) >= 4:
        config_path = sys.argv[3]

    if not dst:
        dst = os.path.splitext(src)[0] + ".docx"

    print("=============================================================================")
    print(f"dbv-md2word v{__version__}, creado por David Bueno Vallejo — libre y gratuito · https://github.com/davidbuenov")
    print("=============================================================================")
    print(f"Analizando {src}...")
    blocks = parse_markdown(src)
    print(f"Generando {dst}...")

    # Configuración por defecto de la CLI tradicional + custom config si se provee
    default_config = load_default_config(config_path)

    create_docx(blocks, dst, default_config, os.path.dirname(os.path.abspath(src)))
    print("¡Proceso completado exitosamente!")

# =============================================================================
# CHANGELOG del parche
# -----------------------------------------------------------------------------
# [Fix]     add_runs_to_paragraph: la detección de enlaces markdown comprobaba
#           '[' , ']', '(' y ')' en cualquier posición del texto (con "in"),
#           no su adyacencia real "](" . Esto provocaba que celdas como
#           "[Sí] Activo (pasivo)" se trataran como enlace y perdieran todo el
#           texto salvo lo que había entre corchetes. Ahora se exige '](' en
#           el texto y se usa un regex anclado (^...$) para extraer texto/URL.
#
# [Feature] Nuevas funciones EMOJI_RANGES / is_emoji_char / _split_run_by_emoji
#           / apply_emoji_font: tras construir el documento (antes de
#           doc.save), se recorren todos los runs (párrafos sueltos y celdas
#           de tabla, incluidas tablas anidadas) y cualquier carácter Unicode
#           de tipo emoji (flechas, dingbats, símbolos misceláneos, bloques
#           suplementarios 1F300-1FAFF, selectores de variación...) se separa
#           en su propio sub-run con fuente "emoji_font" (por defecto
#           "Segoe UI Emoji"), configurable vía config.json. El resto del
#           texto conserva la fuente original del run (body_font, negrita,
#           color, etc.). Esto es necesario porque Word no genera fallback de
#           fuente automático para un run cuyo w:rFonts ya declara
#           explícitamente una fuente sin esos glifos (a diferencia de cuando
#           el usuario escribe/pega el emoji en caliente, donde Word sí
#           reasigna la fuente sobre la marcha).
#
# Para desactivar el cambio de fuente de emoji, basta con poner
# "emoji_font": "" (cadena vacía) en config.json.
# =============================================================================
