# =============================================================================
# dbv-md2word — Conversor personalizable de Markdown a Word (.docx) con interfaz visual local
# Copyright (c) 2026 David Bueno Vallejo · https://github.com/davidbuenov
# Licensed under the MIT License. See LICENSE for details.
# Built with dbv-specs-ops · https://github.com/davidbuenov/dbv-specs-ops
# =============================================================================

import os
import unittest
import tempfile
import docx
from docx import Document

# Añadimos la ruta de la aplicación al path de python
import sys
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from convert_md_to_docx import sanitize_bookmark_name, parse_markdown, create_docx

class TestConverter(unittest.TestCase):
    
    def test_sanitize_bookmark_name(self):
        """Prueba la sanitización de nombres de marcadores en Word."""
        self.assertEqual(sanitize_bookmark_name("Mi Imagen de Prueba"), "MiImagendePrueba")
        self.assertEqual(sanitize_bookmark_name("123 Imagen"), "_123Imagen")
        self.assertEqual(sanitize_bookmark_name("img-con-guiones!_y_mas"), "imgconguiones_y_mas")
        # Límite de longitud
        long_name = "a" * 50
        self.assertEqual(len(sanitize_bookmark_name(long_name)), 40)

    def test_parse_markdown_headings(self):
        """Prueba que los encabezados se parseen correctamente en bloques."""
        # Creamos un archivo markdown temporal
        with tempfile.NamedTemporaryFile(mode='w+', suffix='.md', delete=False, encoding='utf-8') as f:
            f.write("# Título 1\n\n## Subtítulo 2\n\nTexto normal aquí.")
            temp_path = f.name
            
        try:
            blocks = parse_markdown(temp_path)
            self.assertEqual(len(blocks), 3)
            self.assertEqual(blocks[0]['type'], 'heading')
            self.assertEqual(blocks[0]['level'], 1)
            self.assertEqual(blocks[0]['text'], 'Título 1')
            
            self.assertEqual(blocks[1]['type'], 'heading')
            self.assertEqual(blocks[1]['level'], 2)
            self.assertEqual(blocks[1]['text'], 'Subtítulo 2')
            
            self.assertEqual(blocks[2]['type'], 'paragraph')
            self.assertEqual(blocks[2]['text'], 'Texto normal aquí.')
        finally:
            os.remove(temp_path)

    def test_parse_markdown_tables_with_captions(self):
        """Prueba el parseo de tablas con pie o cabecera de tabla asociado."""
        with tempfile.NamedTemporaryFile(mode='w+', suffix='.md', delete=False, encoding='utf-8') as f:
            f.write("Tabla: Mis Datos\n| Col A | Col B |\n|---|---|\n| Val 1 | Val 2 |")
            temp_path = f.name
            
        try:
            blocks = parse_markdown(temp_path)
            self.assertEqual(len(blocks), 1)
            self.assertEqual(blocks[0]['type'], 'table')
            self.assertEqual(blocks[0]['caption'], 'Mis Datos')
            self.assertEqual(len(blocks[0]['rows']), 2)
            self.assertEqual(blocks[0]['rows'][0], ['Col A', 'Col B'])
        finally:
            os.remove(temp_path)

    def test_parse_markdown_images(self):
        """Prueba la detección de imágenes en bloques."""
        with tempfile.NamedTemporaryFile(mode='w+', suffix='.md', delete=False, encoding='utf-8') as f:
            f.write("![Diagrama de flujo](images/diagram.png)")
            temp_path = f.name
            
        try:
            blocks = parse_markdown(temp_path)
            self.assertEqual(len(blocks), 1)
            self.assertEqual(blocks[0]['type'], 'image')
            self.assertEqual(blocks[0]['alt'], 'Diagrama de flujo')
            self.assertEqual(blocks[0]['src'], 'images/diagram.png')
        finally:
            os.remove(temp_path)

    def test_create_docx_generation(self):
        """Prueba que el archivo docx se genere correctamente con los estilos nativos mapeados."""
        blocks = [
            {'type': 'heading', 'level': 1, 'text': 'Título de Prueba'},
            {'type': 'paragraph', 'text': 'Texto que contiene una referencia [Fig. Diagrama de flujo].'},
            {'type': 'image', 'alt': 'Diagrama de flujo', 'src': 'non_existent.png'}
        ]
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            temp_out_path = f.name
            
        try:
            config = {
                'heading_font': 'Arial',
                'body_font': 'Georgia',
                'primary_color': '00FF00',
                'toc_enabled': True
            }
            create_docx(blocks, temp_out_path, config)
            
            # Cargamos el archivo generado para validar que python-docx lo lee
            doc = Document(temp_out_path)
            self.assertTrue(len(doc.paragraphs) > 0)
            
            # Verificamos que se crearon los estilos personalizados
            self.assertIn('codigo', doc.styles)
            self.assertIn('codigo_car', doc.styles)
            
        finally:
            if os.path.exists(temp_out_path):
                os.remove(temp_out_path)

    def test_shift_headings_mapping(self):
        """Prueba que el desplazamiento de títulos (shift_headings) mapee H1 a Title y H2 a Heading 1."""
        blocks = [
            {'type': 'heading', 'level': 1, 'text': 'Título del Documento'},
            {'type': 'heading', 'level': 2, 'text': 'Sección Uno'}
        ]
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            temp_out_path = f.name
            
        try:
            config = {
                'heading_font': 'Arial',
                'body_font': 'Calibri',
                'primary_color': 'FF0000',
                'toc_enabled': False,
                'shift_headings': True
            }
            create_docx(blocks, temp_out_path, config)
            
            doc = Document(temp_out_path)
            # El primer párrafo es el H1 desplazado (Título)
            # El segundo párrafo es el H2 desplazado (Título 1 / Heading 1)
            # Nota: python-docx crea un párrafo vacío por defecto al instanciar Document(), 
            # pero create_docx usa doc.add_paragraph() para los bloques.
            # Verificamos los estilos aplicados
            styles_used = [p.style.name for p in doc.paragraphs if p.text.strip()]
            self.assertEqual(styles_used[0], 'Title')
            self.assertEqual(styles_used[1], 'Heading 1')
            
        finally:
            if os.path.exists(temp_out_path):
                os.remove(temp_out_path)

if __name__ == '__main__':
    unittest.main()
