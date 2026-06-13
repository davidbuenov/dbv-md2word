import re
import os
import sys
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.opc.constants import RELATIONSHIP_TYPE

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink node
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a w:r node
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Add color (Blue)
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0563C1')
    rPr.append(c)

    # Add underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)

    text_node = OxmlElement('w:t')
    text_node.text = text
    new_run.append(text_node)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def add_runs_to_paragraph(p, text):
    pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*|`.*?`|\[.*?\]\(.*?\))')
    parts = pattern.split(text)
    
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            p.add_run(part[2:-2]).bold = True
        elif part.startswith('*') and part.endswith('*'):
            p.add_run(part[1:-1]).italic = True
        elif part.startswith('`') and part.endswith('`'):
            r = p.add_run(part[1:-1])
            r.font.name = 'Courier New'
            r.font.size = Pt(9.5)
            shading = parse_xml(r'<w:shd {} w:fill="F0F0F0"/>'.format(nsdecls('w')))
            r._r.get_or_add_rPr().append(shading)
        elif part.startswith('[') and ']' in part and '(' in part and part.endswith(')'):
            link_text_match = re.search(r'\[(.*?)\]', part)
            link_url_match = re.search(r'\((.*?)\)', part)
            if link_text_match and link_url_match:
                link_text = link_text_match.group(1)
                link_url = link_url_match.group(1)
                add_hyperlink(p, link_text, link_url)
        else:
            p.add_run(part)

def set_cell_background(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    tcPr.append(shd)

def parse_markdown(filepath):
    if not os.path.exists(filepath):
        print(f"Error: El archivo {filepath} no existe.")
        sys.exit(1)
        
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    blocks = []
    current_block = None
    in_code = False
    in_table = False
    
    for line in lines:
        stripped = line.strip()
        
        # Code block handling
        if stripped.startswith('```'):
            if in_code:
                in_code = False
                blocks.append(current_block)
                current_block = None
            else:
                if current_block:
                    blocks.append(current_block)
                in_code = True
                lang = stripped[3:].strip()
                current_block = {'type': 'code', 'lang': lang, 'lines': []}
            continue
            
        if in_code:
            current_block['lines'].append(line.rstrip('\n'))
            continue
            
        # Table handling
        if stripped.startswith('|'):
            if re.match(r'^\|[\s:-|]+$', stripped):
                continue
            row_cells = [cell.strip() for cell in stripped.split('|')[1:-1]]
            if in_table:
                current_block['rows'].append(row_cells)
            else:
                if current_block:
                    blocks.append(current_block)
                in_table = True
                current_block = {'type': 'table', 'rows': [row_cells]}
            continue
        else:
            if in_table:
                in_table = False
                blocks.append(current_block)
                current_block = None
                
        # Empty line
        if not stripped:
            if current_block and current_block['type'] not in ['code', 'table']:
                blocks.append(current_block)
                current_block = None
            continue
            
        # Heading
        heading_match = re.match(r'^(#{1,6})\s+(.*)$', stripped)
        if heading_match:
            if current_block:
                blocks.append(current_block)
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            blocks.append({'type': 'heading', 'level': level, 'text': text})
            current_block = None
            continue
            
        # Bullet list
        bullet_match = re.match(r'^([-*]|\d+\.)\s+(.*)$', stripped)
        indent_match = re.match(r'^(\s+)([-*]|\d+\.)\s+(.*)$', line)
        
        if indent_match:
            indent_spaces = len(indent_match.group(1))
            text = indent_match.group(3)
            level = indent_spaces // 2
            item = {'text': text, 'level': level}
            if current_block and current_block['type'] == 'list':
                current_block['items'].append(item)
            else:
                if current_block:
                    blocks.append(current_block)
                current_block = {'type': 'list', 'items': [item]}
            continue
        elif bullet_match:
            text = bullet_match.group(2)
            item = {'text': text, 'level': 0}
            if current_block and current_block['type'] == 'list':
                current_block['items'].append(item)
            else:
                if current_block:
                    blocks.append(current_block)
                current_block = {'type': 'list', 'items': [item]}
            continue
            
        # Horizontal Rule
        if stripped == '---':
            if current_block:
                blocks.append(current_block)
            blocks.append({'type': 'hr'})
            current_block = None
            continue
            
        # Paragraph
        if current_block and current_block['type'] == 'paragraph':
            current_block['text'] += " " + stripped
        else:
            if current_block:
                blocks.append(current_block)
            current_block = {'type': 'paragraph', 'text': stripped}
            
    if current_block:
        blocks.append(current_block)
        
    return blocks

def create_docx(blocks, output_path):
    doc = Document()
    
    # Configure margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Configure base font style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    first_heading = True
    
    for block in blocks:
        if block['type'] == 'heading':
            level = block['level']
            text = block['text']
            
            p = doc.add_paragraph()
            p.paragraph_format.keep_with_next = True
            run = p.add_run(text)
            run.bold = True
            
            if level == 1:
                run.font.size = Pt(20)
                run.font.color.rgb = RGBColor(31, 78, 121) # Dark Blue
                p.paragraph_format.space_before = Pt(24 if not first_heading else 0)
                p.paragraph_format.space_after = Pt(12)
                first_heading = False
            elif level == 2:
                run.font.size = Pt(15)
                run.font.color.rgb = RGBColor(46, 116, 181) # Medium Blue
                p.paragraph_format.space_before = Pt(18)
                p.paragraph_format.space_after = Pt(8)
            elif level == 3:
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(46, 116, 181)
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
            else:
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(89, 89, 89)
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(4)
                
        elif block['type'] == 'paragraph':
            if not block['text'].strip():
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.15
            add_runs_to_paragraph(p, block['text'])
            
        elif block['type'] == 'list':
            for item in block['items']:
                style_name = 'List Bullet'
                if item['level'] == 1:
                    style_name = 'List Bullet 2'
                elif item['level'] >= 2:
                    style_name = 'List Bullet 3'
                    
                p = doc.add_paragraph(style=style_name)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.line_spacing = 1.15
                add_runs_to_paragraph(p, item['text'])
                
        elif block['type'] == 'hr':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            p_border = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:bottom w:val="single" w:sz="6" w:space="1" w:color="D3D3D3"/></w:pBdr>')
            p._p.get_or_add_pPr().append(p_border)
            
        elif block['type'] == 'code':
            table = doc.add_table(rows=1, cols=1)
            table.autofit = False
            cell = table.cell(0, 0)
            cell.width = Inches(6.5)
            
            set_cell_background(cell, "F5F5F5")
            
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                                  r'<w:top w:val="none"/>'
                                  r'<w:left w:val="single" w:sz="18" w:space="0" w:color="2E74B5"/>'
                                  r'<w:bottom w:val="none"/>'
                                  r'<w:right w:val="none"/>'
                                  r'</w:tcBorders>')
            tcPr.append(tcBorders)
            
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            
            code_text = "\n".join(block['lines'])
            run = p.add_run(code_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(60, 60, 60)
            
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(0)
            spacer.paragraph_format.space_after = Pt(6)
            
        elif block['type'] == 'table':
            rows_data = block['rows']
            if not rows_data:
                continue
                
            cols_count = len(rows_data[0])
            table = doc.add_table(rows=len(rows_data), cols=cols_count)
            table.style = 'Table Grid'
            table.autofit = True
            
            for r_idx, row_data in enumerate(rows_data):
                row = table.rows[r_idx]
                is_header = (r_idx == 0)
                
                for c_idx, cell_value in enumerate(row_data):
                    if c_idx >= len(row.cells):
                        continue
                    cell = row.cells[c_idx]
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.space_before = Pt(2)
                    
                    if is_header:
                        set_cell_background(cell, "2E74B5")
                        run = p.add_run()
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        add_runs_to_paragraph(p, cell_value)
                    else:
                        if r_idx % 2 == 0:
                            set_cell_background(cell, "F2F5F8")
                        add_runs_to_paragraph(p, cell_value)
            
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(0)
            spacer.paragraph_format.space_after = Pt(8)
            
    doc.save(output_path)

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Uso: python convert_md_to_docx.py <archivo.md> [archivo.docx]")
        sys.exit(1)
        
    src = sys.argv[1]
    if len(sys.argv) >= 3:
        dst = sys.argv[2]
    else:
        dst = os.path.splitext(src)[0] + ".docx"
        
    print(f"Analizando {src}...")
    blocks = parse_markdown(src)
    print(f"Generando {dst}...")
    create_docx(blocks, dst)
    print("¡Proceso completado exitosamente!")
